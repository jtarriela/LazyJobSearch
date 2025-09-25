"""
Enhanced Resume Parser with comprehensive LLM extraction.
Properly integrates with OpenAI/Gemini for complete resume parsing.
"""

from __future__ import annotations
import logging
import re
import asyncio
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Union
import os

logger = logging.getLogger(__name__)

@dataclass
class ResumeSection:
    """Represents a parsed section of a resume"""
    name: str
    content: str
    start_index: int
    end_index: int

@dataclass
class ParsedResume:
    """Enhanced parsed resume with comprehensive data"""
    fulltext: str
    sections: Dict[str, str]
    skills: List[str]  # Complete list of all skills
    years_of_experience: Optional[float]
    education_level: Optional[str]
    contact_info: Dict[str, str]
    word_count: int
    char_count: int
    
    # Enhanced fields from LLM parsing
    full_name: Optional[str] = None
    experience: List[Dict[str, str]] = None
    education: List[Dict[str, str]] = None
    certifications: List[str] = None
    summary: Optional[str] = None
    awards: List[str] = None
    publications: List[str] = None
    projects: List[Dict[str, str]] = None
    clearance_level: Optional[str] = None
    skills_structured: Dict[str, List[str]] = None
    parsing_method: str = "llm"
    
    def __post_init__(self):
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.certifications is None:
            self.certifications = []
        if self.awards is None:
            self.awards = []
        if self.publications is None:
            self.publications = []
        if self.projects is None:
            self.projects = []
        if self.skills_structured is None:
            self.skills_structured = {}

class ResumeParseError(Exception):
    """Custom exception for resume parsing errors"""
    pass

class ResumeParser:
    """Enhanced parser for comprehensive resume extraction"""
    
    def __init__(self, use_llm: bool = True):
        """Initialize the resume parser with LLM support
        
        Args:
            use_llm: Whether to use LLM for parsing (should always be True for production)
        """
        self.use_llm = use_llm
        
        if self.use_llm:
            from libs.resume.llm_service import create_llm_service, LLMConfig, LLMProvider
            
            # Create LLM service with configured provider
            config = LLMConfig()
            # Convert string provider to enum
            if config.provider == "openai":
                provider = LLMProvider.OPENAI
            elif config.provider == "gemini":
                provider = LLMProvider.GEMINI
            else:
                raise ValueError(f"Unsupported provider: {config.provider}")
                
            self.llm_service = create_llm_service(provider=provider)
            logger.info(f"Initialized LLM parser with {config.provider} provider")
    
    def parse_file(self, file_path: Union[str, Path]) -> ParsedResume:
        """Parse a resume file with comprehensive extraction
        
        Args:
            file_path: Path to resume file (PDF, DOCX, TXT)
            
        Returns:
            ParsedResume with all extracted content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise ResumeParseError(f"File not found: {file_path}")
        
        try:
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                fulltext = self._extract_pdf_text(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                fulltext = self._extract_docx_text(file_path)
            elif file_path.suffix.lower() in ['.txt']:
                fulltext = file_path.read_text(encoding='utf-8')
            else:
                raise ResumeParseError(f"Unsupported file type: {file_path.suffix}")
            
            if not fulltext.strip():
                raise ResumeParseError("No text content extracted from file")
            
            # Parse with LLM for comprehensive extraction
            return self._parse_text(fulltext)
            
        except Exception as e:
            logger.error(f"Failed to parse resume {file_path}: {e}")
            raise ResumeParseError(f"Parsing failed: {str(e)}") from e
    
    def parse_text(self, text: str) -> ParsedResume:
        """Parse resume text directly"""
        return self._parse_text(text)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                # Try alternative extraction with PyPDF2 as fallback
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                except:
                    pass
            
            logger.debug(f"Extracted {len(text)} characters from PDF")
            return text.strip()
            
        except ImportError:
            raise ResumeParseError("PDF parsing libraries not available. Install with: pip install pdfplumber PyPDF2")
        except Exception as e:
            raise ResumeParseError(f"PDF extraction failed: {str(e)}") from e
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            document = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.debug(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except ImportError:
            raise ResumeParseError("DOCX parsing library not available. Install with: pip install python-docx")
        except Exception as e:
            raise ResumeParseError(f"DOCX extraction failed: {str(e)}") from e
    
    def _parse_text(self, text: str) -> ParsedResume:
        """Parse resume text using LLM for comprehensive extraction"""
        
        if not self.use_llm:
            raise ValueError("LLM parsing is required for comprehensive extraction")
        
        # Clean text
        normalized_text = self._normalize_text(text)
        
        # Parse with LLM
        return asyncio.run(self._parse_with_llm(normalized_text))
    
    async def _parse_with_llm(self, text: str) -> ParsedResume:
        """Parse resume using enhanced LLM service"""
        logger.info("Starting comprehensive LLM parsing...")
        
        # Get structured data from LLM
        llm_data, responses = await self.llm_service.parse_resume(text)
        
        # Extract sections for backward compatibility
        sections = self._extract_sections_from_text(text)
        
        # Build comprehensive contact info
        contact_info = {}
        if llm_data.email:
            contact_info['email'] = llm_data.email
        if llm_data.phone:
            contact_info['phone'] = llm_data.phone
        if llm_data.location:
            contact_info['location'] = llm_data.location
        if llm_data.links:
            if llm_data.links.linkedin:
                contact_info['linkedin'] = llm_data.links.linkedin
            if llm_data.links.github:
                contact_info['github'] = llm_data.links.github
            if llm_data.links.portfolio:
                contact_info['portfolio'] = llm_data.links.portfolio
        
        # Convert experience to dict format
        experience_list = []
        for exp in llm_data.experience:
            exp_dict = {
                'title': exp.title,
                'company': exp.company,
                'duration': exp.duration,
                'location': exp.location,
                'description': exp.description,
                'bullets': exp.bullets,
                'technologies': exp.technologies
            }
            experience_list.append(exp_dict)
        
        # Convert education to dict format  
        education_list = []
        for edu in llm_data.education:
            edu_dict = {
                'degree': edu.degree,
                'field': edu.field,
                'institution': edu.institution,
                'year': edu.year,
                'gpa': edu.gpa,
                'specialization': edu.specialization
            }
            education_list.append(edu_dict)
        
        # Convert projects to dict format
        projects_list = []
        for proj in llm_data.projects:
            proj_dict = {
                'name': proj.name,
                'description': proj.description,
                'technologies': proj.technologies,
                'link': proj.link
            }
            projects_list.append(proj_dict)
        
        # Extract certifications as strings
        cert_list = []
        for cert in llm_data.certifications:
            if isinstance(cert, str):
                cert_list.append(cert)
            else:
                cert_name = cert.name
                if cert.issuer:
                    cert_name += f" ({cert.issuer})"
                if cert.date:
                    cert_name += f" - {cert.date}"
                cert_list.append(cert_name)
        
        # Build structured skills dict
        skills_structured = {}
        if llm_data.skills_structured:
            skills_structured = {
                'programming': llm_data.skills_structured.programming,
                'frameworks': llm_data.skills_structured.frameworks,
                'tools': llm_data.skills_structured.tools,
                'databases': llm_data.skills_structured.databases,
                'cloud': llm_data.skills_structured.cloud,
                'ml_ai': llm_data.skills_structured.ml_ai,
                'other': llm_data.skills_structured.other,
                'all': llm_data.skills_structured.all
            }
        
        # Log extraction results
        logger.info(f"LLM extraction complete:")
        logger.info(f"  - Name: {llm_data.full_name}")
        logger.info(f"  - Skills: {len(llm_data.skills)} total")
        logger.info(f"  - Experience: {len(experience_list)} positions")
        logger.info(f"  - Education: {len(education_list)} degrees")
        logger.info(f"  - Certifications: {len(cert_list)}")
        logger.info(f"  - Awards: {len(llm_data.awards)}")
        logger.info(f"  - Publications: {len(llm_data.publications)}")
        logger.info(f"  - Projects: {len(projects_list)}")
        
        if len(llm_data.skills) < 10 and "python" in text.lower():
            logger.warning("Low skill count detected - may need to re-run extraction")
        
        return ParsedResume(
            fulltext=text,
            sections=sections,
            skills=llm_data.skills or [],
            years_of_experience=llm_data.years_of_experience,
            education_level=llm_data.education_level,
            contact_info=contact_info,
            word_count=len(text.split()),
            char_count=len(text),
            full_name=llm_data.full_name,
            experience=experience_list,
            education=education_list,
            certifications=cert_list,
            summary=llm_data.summary,
            awards=llm_data.awards,
            publications=llm_data.publications,
            projects=projects_list,
            clearance_level=llm_data.clearance_level,
            skills_structured=skills_structured,
            parsing_method="llm"
        )
    
    def _normalize_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()
    
    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """Extract major sections from text for backward compatibility"""
        sections = {}
        text_lower = text.lower()
        
        # Common section headers
        section_patterns = {
            'education': r'(?i)\b(education|academic)\b',
            'experience': r'(?i)\b(experience|employment|work history)\b',
            'skills': r'(?i)\b(skills|technical skills|competencies)\b',
            'projects': r'(?i)\b(projects|portfolio)\b',
            'certifications': r'(?i)\b(certifications?|licenses?)\b',
            'awards': r'(?i)\b(awards?|honors?|achievements?)\b',
            'publications': r'(?i)\b(publications?|papers?)\b'
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text)
            if match:
                start = match.start()
                # Find next section or end of text
                next_section_start = len(text)
                for other_pattern in section_patterns.values():
                    if other_pattern == pattern:
                        continue
                    other_matches = re.finditer(other_pattern, text[start + 1:])
                    for other_match in other_matches:
                        pos = start + 1 + other_match.start()
                        if pos < next_section_start:
                            next_section_start = pos
                            break
                
                section_content = text[start:next_section_start].strip()
                if section_content:
                    sections[section_name] = section_content
        
        return sections

def create_resume_parser(use_llm: bool = True) -> ResumeParser:
    """Factory function to create enhanced resume parser"""
    if not use_llm:
        raise ValueError("LLM parsing is required for comprehensive extraction")
    return ResumeParser(use_llm=True)