"""Resume file parsers for different formats."""

import re
from pathlib import Path
from typing import Dict, Any
from abc import ABC, abstractmethod

import PyPDF2
from docx import Document


class ResumeParser(ABC):
    """Base class for resume file parsers."""
    
    @abstractmethod
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse a resume file and return structured content.
        
        Returns:
            Dict with 'fulltext' and 'sections' keys
        """
        pass
    
    def _extract_sections(self, text: str) -> Dict[str, Any]:
        """Extract structured sections from resume text."""
        sections = {}
        
        # Common resume section patterns
        section_patterns = {
            'summary': r'(?:SUMMARY|PROFESSIONAL SUMMARY|CAREER SUMMARY|OBJECTIVE)[\n\r]+(.*?)(?=\n[A-Z\s]{2,}|$)',
            'experience': r'(?:EXPERIENCE|WORK EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT)[\n\r]+(.*?)(?=\n[A-Z\s]{2,}|$)',
            'education': r'(?:EDUCATION|ACADEMIC BACKGROUND|QUALIFICATIONS)[\n\r]+(.*?)(?=\n[A-Z\s]{2,}|$)',
            'skills': r'(?:SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|TECHNOLOGIES)[\n\r]+(.*?)(?=\n[A-Z\s]{2,}|$)',
            'projects': r'(?:PROJECTS|KEY PROJECTS|NOTABLE PROJECTS)[\n\r]+(.*?)(?=\n[A-Z\s]{2,}|$)',
        }
        
        for section_name, pattern in section_patterns.items():
            matches = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                sections[section_name] = matches.group(1).strip()
        
        return {
            'sections': sections,
            'section_count': len(sections)
        }


class PDFParser(ResumeParser):
    """Parser for PDF resume files."""
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse a PDF resume file."""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
            
            # Join all pages
            fulltext = '\n\n'.join(text_content)
            
            # Clean up the text
            fulltext = self._clean_text(fulltext)
            
            # Extract sections
            sections = self._extract_sections(fulltext)
            
            return {
                'fulltext': fulltext,
                'sections': sections,
                'page_count': len(pdf_reader.pages),
                'file_type': 'pdf'
            }
            
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean up extracted PDF text."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Add space after punctuation
        
        return text.strip()


class DocxParser(ResumeParser):
    """Parser for Word document resume files."""
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse a DOCX resume file."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + tables_text
            fulltext = '\n\n'.join(all_text)
            
            # Clean up the text
            fulltext = self._clean_text(fulltext)
            
            # Extract sections
            sections = self._extract_sections(fulltext)
            
            return {
                'fulltext': fulltext,
                'sections': sections,
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'file_type': 'docx'
            }
            
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean up extracted DOCX text."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page breaks and other formatting artifacts
        text = re.sub(r'\f', '\n', text)  # Form feed to newline
        text = re.sub(r'\x0b', '\n', text)  # Vertical tab to newline
        
        return text.strip()


class TextParser(ResumeParser):
    """Parser for plain text resume files."""
    
    def parse(self, file_path: Path) -> Dict[str, Any]:
        """Parse a plain text resume file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                fulltext = file.read()
            
            # Basic cleaning
            fulltext = self._clean_text(fulltext)
            
            # Extract sections
            sections = self._extract_sections(fulltext)
            
            return {
                'fulltext': fulltext,
                'sections': sections,
                'file_type': 'txt'
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    fulltext = file.read()
                
                fulltext = self._clean_text(fulltext)
                sections = self._extract_sections(fulltext)
                
                return {
                    'fulltext': fulltext,
                    'sections': sections,
                    'file_type': 'txt',
                    'encoding': 'latin-1'
                }
            except Exception as e:
                raise ValueError(f"Error parsing text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean up text content."""
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        return text.strip()